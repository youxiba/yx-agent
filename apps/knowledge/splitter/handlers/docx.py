# coding=utf-8
from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.oxml.ns import qn
from ..file import File
from ..spi import BaseSplitHandle, ParagraphRaw, SplitOptions
from ..split_model import SplitModel


def _iter_block_items(doc):
    """按文档流顺序迭代段落与表格"""
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, doc)


class DocSplitHandle(BaseSplitHandle):
    def support(self, file: File) -> bool:
        return file.suffix == ".docx"

    def get_content(self, file: File, save_image: bool = False) -> str:
        doc = DocxDocument(file.path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def handle(self, file: File, opts: SplitOptions) -> list[ParagraphRaw]:
        doc = DocxDocument(file.path)
        sections: list[tuple[str, list[str]]] = [("", [])]
        for block in _iter_block_items(doc):
            if isinstance(block, DocxParagraph):
                text = block.text.strip()
                if not text:
                    continue
                style = (block.style.name or "").lower()
                if style.startswith("heading") or style == "title":
                    sections.append((text, []))
                else:
                    sections[-1][1].append(text)
            else:  # 表格转 MD
                sections[-1][1].append(self._table_to_md(block))
        sm = SplitModel(opts.content_level_pattern, opts.limit, opts.with_filter)
        raws: list[ParagraphRaw] = []
        for title, lines in sections:
            raws.extend(sm.parse_to_paragraphs("\n".join(lines), title=title))
        return raws

    @staticmethod
    def _table_to_md(table: DocxTable) -> str:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            return ""
        md = "| " + " | ".join(rows[0]) + " |\n| " + " | ".join("---" for _ in rows[0]) + " |\n"
        for r in rows[1:]:
            md += "| " + " | ".join(r) + " |\n"
        return md