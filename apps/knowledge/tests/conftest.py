# coding=utf-8
import io
import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook
from reportlab.pdfgen import canvas
from pypdf import PdfWriter
from identity.models import User, Role, Workspace, WorkspaceMember


@pytest.fixture
def workspace():
    u = User.objects.create_user(username="kw_owner", email="kw@x.cn", password="Passw0rd!", role=Role.ADMIN)
    ws = Workspace.objects.create(name="测试空间", owner=u)
    WorkspaceMember.objects.create(workspace=ws, user=u, role=Role.WORKSPACE_MANAGE)
    return ws


@pytest.fixture
def user(workspace):
    return workspace.owner


@pytest.fixture
def knowledge(workspace, user):
    from knowledge.models import Knowledge
    return Knowledge.objects.create(name="集成库", workspace=workspace, user=user, embedding_model_id="fake-embed")


@pytest.fixture
def document(knowledge):
    from knowledge.models import Document
    return Document.objects.create(knowledge=knowledge, name="样本.md", meta={"file_path": ""})


@pytest.fixture
def make_paragraph(document):
    from knowledge.models import Paragraph
    def _mk(doc, content, title="标题"):
        return Paragraph.objects.create(document=doc, knowledge=doc.knowledge, title=title, content=content)
    return _mk


@pytest.fixture
def fake_gateway(monkeypatch):
    """桩 ModelGateway：embed 返回确定性向量"""
    class FakeEmbed:
        def embed_query(self, text):
            return [1.0] + [0.0] * 1023
        def embed_documents(self, texts):
            return [[0.1] * 1024 for _ in texts]
        def get_vector_dim(self):
            return 1024
    class FakeGateway:
        def get_model(self, model_id):
            return FakeEmbed()
    monkeypatch.setattr("knowledge.services.embedder.gateway", FakeGateway())
    return FakeGateway()


# ---- 样本文件生成器 ----
def make_md_bytes():
    return "# 标题一\n内容 A 的正文。\n\n## 子标题\n内容 B。\n".encode("utf-8")


def make_docx_bytes():
    doc = DocxDocument()
    doc.add_heading("产品介绍", level=1)
    doc.add_paragraph("这是正文第一段。")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_xlsx_bytes():
    wb = Workbook()
    ws = wb.active
    ws.title = "人员"
    ws.append(["姓名", "部门"])
    ws.append(["张三", "研发"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def make_pdf_bytes():
    raw = io.BytesIO()
    c = canvas.Canvas(raw)
    c.drawString(72, 720, "Hello MaxKB PDF")
    c.save()
    raw.seek(0)
    writer = PdfWriter()
    writer.append(raw)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()